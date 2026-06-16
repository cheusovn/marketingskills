#!/usr/bin/env python3
import os
import re
from pathlib import Path
from PyPDF2 import PdfReader
from pdf2image import convert_from_path
import pytesseract
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# Configuration
UPLOAD_DIR = '/root/.claude/uploads/3e540a7b-5ab9-5cc8-9677-e3d80f77fa1e'
OUTPUT_DIR = '/home/user/marketingskills'

# Replacement patterns
REPLACEMENTS = [
    (r'ИП Чеусова Николая ИГоревича', 'ИП Пугачеву Ольгу Николаевну'),
    (r'ИП Чеусова Николая Игоревича', 'ИП Пугачеву Ольгу Николаевну'),
    (r'Чеусова Николая ИГоревича', 'Пугачеву Ольгу Николаевну'),
    (r'Чеусова Николая Игоревича', 'Пугачеву Ольгу Николаевну'),
    (r'Чеусова', 'Пугачева'),
    (r'Николай Игоревич', 'Ольга Николаевна'),
    (r'Николай ИГоревич', 'Ольга Николаевна'),
]

def extract_text_direct(pdf_path):
    """Try to extract text directly from PDF."""
    text = ""
    try:
        pdf_reader = PdfReader(pdf_path)
        for page in pdf_reader.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + '\n'
    except Exception as e:
        print(f"  Warning: Direct extraction failed: {e}")
    return text

def extract_text_ocr(pdf_path):
    """Extract text from PDF using OCR."""
    print(f"  Using OCR to extract text...")
    text = ""
    try:
        images = convert_from_path(pdf_path)
        for page_num, image in enumerate(images, 1):
            print(f"    Processing page {page_num}...")
            page_text = pytesseract.image_to_string(image, lang='rus+eng')
            if page_text:
                text += page_text + '\n'
    except Exception as e:
        print(f"  OCR extraction failed: {e}")
    return text

def extract_text_from_pdf(pdf_path):
    """Extract text from PDF file, trying direct extraction first, then OCR."""
    print(f"Extracting text from: {os.path.basename(pdf_path)}")

    # Try direct extraction first
    text = extract_text_direct(pdf_path)

    # If direct extraction failed or returned minimal text, use OCR
    if len(text.strip()) < 100:
        print(f"  Direct extraction insufficient ({len(text)} chars), trying OCR...")
        text = extract_text_ocr(pdf_path)

    if text.strip():
        print(f"✓ Extracted {len(text)} characters")
    else:
        print(f"⚠ No text could be extracted")

    return text

def replace_names(text):
    """Replace all instances of the names."""
    for old, new in REPLACEMENTS:
        text = re.sub(old, new, text)
    return text

def create_word_document(text, output_path, preserve_structure=True):
    """Create a Word document from extracted text."""
    print(f"Creating Word document...")

    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # Add content preserving structure
    if preserve_structure:
        # Split by double newlines to preserve paragraphs
        blocks = text.split('\n\n')

        for block in blocks:
            lines = block.split('\n')
            block_text = ' '.join(line.strip() for line in lines if line.strip())

            if block_text:
                p = doc.add_paragraph(block_text)
                p.paragraph_format.space_after = Pt(12)
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.line_spacing = 1.15
    else:
        # Add line by line
        for line in text.split('\n'):
            if line.strip():
                p = doc.add_paragraph(line.strip())
                p.paragraph_format.space_after = Pt(6)
            else:
                doc.add_paragraph()

    # Save
    try:
        doc.save(output_path)
        print(f"✓ Saved: {os.path.basename(output_path)}")
        return True
    except Exception as e:
        print(f"Error saving document: {e}")
        return False

def main():
    """Main processing function."""
    print("=" * 70)
    print("Document Processing Script - PDF to DOCX Converter")
    print("=" * 70)

    # Find PDF files
    pdf_files = sorted([f for f in os.listdir(UPLOAD_DIR) if f.endswith('.pdf')])

    if not pdf_files:
        print(f"Error: No PDF files found in {UPLOAD_DIR}")
        return

    print(f"\nFound {len(pdf_files)} PDF file(s):")
    for i, f in enumerate(pdf_files, 1):
        size = os.path.getsize(os.path.join(UPLOAD_DIR, f)) / 1024
        print(f"  {i}. {f} ({size:.1f} KB)")

    results = []

    # Process each PDF
    for i, pdf_file in enumerate(pdf_files, 1):
        pdf_path = os.path.join(UPLOAD_DIR, pdf_file)

        print(f"\n{'=' * 70}")
        print(f"Processing file {i}/{len(pdf_files)}: {pdf_file}")
        print('=' * 70)

        # Extract text
        text = extract_text_from_pdf(pdf_path)
        if not text.strip():
            print(f"Skipping {pdf_file} - no text could be extracted")
            continue

        # Replace names
        print("Replacing names and requisites...")
        replaced_text = replace_names(text)

        # Determine output filename
        if 'ekvizit' in pdf_file.lower():
            output_name = 'Реквизиты_Пугачева.docx'
        else:
            # Extract meaningful name from first line
            first_line = replaced_text.split('\n')[0][:40]
            output_name = f'Договор_Пугачева.docx'

        output_path = os.path.join(OUTPUT_DIR, output_name)

        # Remove old file if exists
        if os.path.exists(output_path):
            os.remove(output_path)

        # Create Word document
        if create_word_document(replaced_text, output_path, preserve_structure=True):
            results.append(output_path)

    # Summary
    print("\n" + "=" * 70)
    if results:
        print(f"✅ Processing complete!")
        print(f"Created {len(results)} document(s):\n")
        for doc in results:
            size = os.path.getsize(doc) / 1024
            print(f"  ✓ {os.path.basename(doc)} ({size:.1f} KB)")
    else:
        print("⚠ No documents were created")
    print("=" * 70)

if __name__ == '__main__':
    main()
