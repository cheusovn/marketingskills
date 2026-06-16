#!/usr/bin/env python3
import os
import re
from pathlib import Path
from PyPDF2 import PdfReader
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# Configuration
UPLOAD_DIR = '/root/.claude/uploads/3e540a7b-5ab9-5cc8-9677-e3d80f77fa1e'
OUTPUT_DIR = '/home/user/marketingskills'

# Replacement patterns
REPLACEMENTS = [
    (r'ИП Чеусова Николая ИГоревича', 'ИП Пугачеву Ольгу Николаевну'),
    (r'ИП Чеусова Николая Игоревича', 'ИП Пугачеву Ольгу Николаевну'),
    (r'Чеусова Николая ИГоревича', 'Пугачеву Ольгу Николаевну'),
    (r'Чеусова Николая Игоревича', 'Пугачеву Ольгу Николаевну'),
    (r'Чеусова', 'Пугачева'),
    (r'Николай Игоревич', 'Ольга Николаевна'),
    (r'Николай ИГоревич', 'Ольга Николаевна'),
]

def extract_text_from_pdf(pdf_path):
    """Extract text from PDF file."""
    print(f"Extracting text from: {pdf_path}")
    text = ""

    try:
        pdf_reader = PdfReader(pdf_path)
        num_pages = len(pdf_reader.pages)

        for page_num, page in enumerate(pdf_reader.pages, 1):
            page_text = page.extract_text()
            if page_text:
                text += page_text + '\n'

        print(f"✓ Extracted {len(text)} characters from {num_pages} pages")
    except Exception as e:
        print(f"Error extracting PDF: {e}")
        return ""

    return text

def replace_names(text):
    """Replace all instances of the names."""
    for old, new in REPLACEMENTS:
        text = re.sub(old, new, text, flags=re.IGNORECASE)
    return text

def create_word_document(text, output_path):
    """Create a Word document from extracted text."""
    print(f"Creating Word document: {output_path}")

    doc = Document()

    # Add content
    for line in text.split('\n'):
        if line.strip():
            p = doc.add_paragraph(line.strip())
            p.paragraph_format.space_after = Pt(6)
        else:
            doc.add_paragraph()

    # Save
    try:
        doc.save(output_path)
        print(f"✓ Saved: {output_path}")
        return True
    except Exception as e:
        print(f"Error saving document: {e}")
        return False

def main():
    """Main processing function."""
    print("=" * 60)
    print("Document Processing Script")
    print("=" * 60)

    # Find PDF files
    pdf_files = sorted([f for f in os.listdir(UPLOAD_DIR) if f.endswith('.pdf')])

    if len(pdf_files) < 2:
        print(f"Error: Expected at least 2 PDF files, found {len(pdf_files)}")
        return

    print(f"\nFound {len(pdf_files)} PDF files:")
    for i, f in enumerate(pdf_files, 1):
        print(f"  {i}. {f}")

    results = []

    # Process each PDF
    for i, pdf_file in enumerate(pdf_files, 1):
        pdf_path = os.path.join(UPLOAD_DIR, pdf_file)

        print(f"\n--- Processing file {i}/{len(pdf_files)} ---")

        # Extract text
        text = extract_text_from_pdf(pdf_path)
        if not text:
            print(f"Skipping {pdf_file} - no text extracted")
            continue

        # Replace names
        print("Replacing names...")
        replaced_text = replace_names(text)

        # Determine output filename
        if 'ekvizit' in pdf_file.lower():
            output_name = 'Реквизиты_Пугачева.docx'
        else:
            output_name = f'Документ_{i}_Пугачева.docx'

        output_path = os.path.join(OUTPUT_DIR, output_name)

        # Create Word document
        if create_word_document(replaced_text, output_path):
            results.append(output_path)

    # Summary
    print("\n" + "=" * 60)
    print(f"✅ Processing complete!")
    print(f"Created {len(results)} document(s):")
    for doc in results:
        print(f"  ✓ {os.path.basename(doc)}")
    print("=" * 60)

if __name__ == '__main__':
    main()
